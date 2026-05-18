"""
生成两版简历:
  A - 全面翻新（面向AI产品经理）
  B - 加产品项目 + 微调
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_DIR = os.path.expanduser("~/Desktop")

def set_cell_border(cell, **kwargs):
    """Set cell border."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            for attr, val in kwargs[edge].items():
                element.set(qn(f'w:{attr}'), str(val))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_section_heading(doc, text):
    """Add a section heading with a bottom border line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_item(doc, left_text, right_text="", left_bold=False, left_size=11, right_size=10):
    """Add a line with left (bold optional) and right-aligned text (e.g. date)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if left_text:
        run_l = p.add_run(left_text)
        run_l.bold = left_bold
        run_l.font.size = Pt(left_size)
        run_l.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    if right_text:
        # Add tab stop at right margin
        tab_stops = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), '9072')  # Right margin for US letter 1" margins
        tab_stops.append(tab)
        p.paragraph_format.tab_stops._element.append(tab)
        # Actually python-docx doesn't expose tab_stops easily. Just use spaces or add as separate run.
        run_r = p.add_run(f"    {right_text}")
        run_r.font.size = Pt(right_size)
        run_r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    return p

def add_bullet(doc, text, font_size=10, indent_level=0):
    """Add a bullet point."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5 + indent_level * 0.5)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    run = p.add_run("· " + text)
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p

def add_body(doc, text, font_size=10, color=None):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p

def set_default_style(doc):
    """Set default font for the document."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_name_header(doc, name):
    """Add centered name."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

def add_contact_line(doc, contact_str):
    """Add centered contact info."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(contact_str)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ═══════════════════════════════════════════
# VERSION A — 全面翻新（AI产品经理）
# ═══════════════════════════════════════════
def build_version_a():
    doc = Document()
    set_default_style(doc)

    # Set page margins
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Header
    add_name_header(doc, "夏樱溪")
    add_contact_line(doc, "178 0265 8725  |  1245987829@qq.com  |  女 · 汉族 · 2001.01")

    # ═══ 自我评价 ═══
    add_section_heading(doc, "自我评价")

    self_eval = (
        "零编程基础3天从不会翻墙到独立交付全功能App并提交 App Store 审核，用行动证明了快速学习和从0到1的落地能力。"
        "硕士论文依靠 AI 工具完成8倍常规工作量实证分析并获优秀评分。"
        "信永中和2个IPO+5个年审项目跨部门协调经验，善于拆解复杂流程、核算投入产出。"
        "辩论队带队校赛第三、国护队首批双考核通过，抗压自驱、沟通推动是底色。"
        "对各领域知识保持好奇并能融会贯通迁移应用。"
    )
    add_body(doc, self_eval, font_size=10)

    # ═══ 产品项目 ═══
    add_section_heading(doc, "产品项目")
    add_item(doc, "土豆ToDo — AI辅助独立开发习惯打卡App", "2025.05", left_bold=True)
    add_bullet(doc, "零编程基础自学 vibe coding，3天内从不会翻墙到完成全功能产品，已提交 App Store 审核")
    add_bullet(doc, "全程使用 Cursor / Claude 辅助开发，涵盖 SwiftUI 前端、SwiftData 数据层、Widget 小组件")
    add_bullet(doc, "独立完成产品定位、交互设计、数据模型、App Groups 共享架构，理解 prompt engineering 在真实产品开发中的边界与效率")

    # ═══ 实习经历 ═══
    add_section_heading(doc, "实习经历")

    add_item(doc, "信永中和审计事务所 / 审计实习生", "2023.11 — 2024.06", left_bold=True)
    add_bullet(doc, "参与 2 个 IPO 项目与 5 个年审项目，跨项目组协调对接甲方财务、律师、评估等多部门")
    add_bullet(doc, "完成穿行测试 140 项、签收单稽核 3,287 份，固定资产抽盘两百余项，审查高管银行流水 2,000+ 条")
    add_bullet(doc, "编纂审计底稿 26 份，稽核合同关键条款 508 份，上级满意度 99%")

    add_item(doc, "国家税务局徐州市税务第一分局 / 暑期实习生", "2020.07 — 2020.09", left_bold=True)
    add_bullet(doc, "承包契税全部门文印工作，发现流程瓶颈并改进方法，文印效率提升近 3 倍")
    add_bullet(doc, "纳税高峰期为群众答疑引导、维持秩序，满意度超 90%，同批唯一优秀实习生")

    # ═══ 校园经历 ═══
    add_section_heading(doc, "校园经历")

    add_item(doc, "西安理工大学国旗护卫队 / 刀手", "2022.09 — 2024.09", left_bold=True)
    add_bullet(doc, "首批通过擎旗手与指挥刀手双项资格考核，参与校运会、换旗仪式等活动 13 次，零失误")

    add_item(doc, "经济管理学院 / 团支书", "2018.09 — 2022.06", left_bold=True)
    add_bullet(doc, "组织团支部活动 30+ 场；策划门票制冰雪节晚会，带领各团支书完成宣传曲制作、8 轮线下宣讲，参与人次过百")
    add_bullet(doc, "身兼班长、团支书、学习委员三职，对接辅导员与各科老师，处理文件 800+ 份")

    # ═══ 教育背景 ═══
    add_section_heading(doc, "教育背景")
    add_item(doc, "西安理工大学 / 会计专业硕士", "2022.09 — 2025.06")
    add_item(doc, "南京财经大学红山学院 / 金融专业本科", "2018.09 — 2022.06")

    # ═══ 奖项荣誉 ═══
    add_section_heading(doc, "奖项荣誉")
    awards = [
        "2024  国旗卫士、优秀工作者、国家学业奖学金",
        "2023  校运动会团体及个人一等奖、二等学业奖学金",
        "2020  徐州市「彭聚菁英」暑期实习优秀实习生",
        "2019  校祖国 70 周年主题演讲大赛第 2 名、优秀班干部",
        "2018  校辩论赛第 3 名、优秀班干部、优秀团员、二等奖学金",
    ]
    for a in awards:
        add_bullet(doc, a)

    path = os.path.join(OUT_DIR, "夏樱溪-简历-AI产品经理版A.docx")
    doc.save(path)
    print(f"Version A saved: {path}")


# ═══════════════════════════════════════════
# VERSION B — 加产品项目 + 微调
# ═══════════════════════════════════════════
def build_version_b():
    doc = Document()
    set_default_style(doc)

    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Header
    add_name_header(doc, "夏樱溪")
    add_contact_line(doc, "178 0265 8725  |  1245987829@qq.com  |  女 · 汉族 · 2001.01")

    # ═══ 产品项目 (新增) ═══
    add_section_heading(doc, "产品项目")
    add_item(doc, "土豆ToDo — AI辅助独立开发习惯打卡App", "2025.05", left_bold=True)
    add_bullet(doc, "零编程基础自学 vibe coding，3天内从不会翻墙到完成全功能产品，已提交 App Store 审核")
    add_bullet(doc, "全程使用 Cursor / Claude 辅助开发（SwiftUI + SwiftData + Widget），独立完成产品设计与数据架构")

    # ═══ 岗位胜任力 ═══
    add_section_heading(doc, "岗位胜任力")

    comp_b = [
        ("AI工具热情与熟练使用",
         "硕士论文用 DeepSeek、GPT、Cursor 辅助完成 8 倍常规工作量的实证代码，盲审获评优秀论文。独立用 vibe coding 完成习惯打卡 App 全流程开发。"),
        ("快速学习与跨界迁移",
         "一周精通高端 PPT 制作；内容创作单条破 400 赞；AI P图活动获第 4 名及 500 元奖金。"),
        ("逻辑拆解与需求分析",
         "会计专业训练了拆解复杂流程、核算成本收益的能力，直接迁移至产品功能决策与用户路径分析。"),
        ("沟通协作与跨团队推动",
         "信永中和多项目跨部门协调经验；辩论队带队校赛第 3 名、演讲比赛第 2 名、职业规划大赛院第 1 名。"),
        ("抗压自驱与细节落地",
         "国护队首批通过刀手+擎旗手双考核，晨训出勤率 100%，参与仪仗活动 13 次零失误。税务局实习期间主动优化流程，文印效率提升近 3 倍，同批唯一优秀实习生。"),
    ]
    for i, (title, body) in enumerate(comp_b):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        run_t = p.add_run(f"{i+1}、{title}")
        run_t.bold = True
        run_t.font.size = Pt(10)
        run_t.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        add_body(doc, body)

    # ═══ 校园及实践经历 ═══
    add_section_heading(doc, "校园及实践经历")

    add_item(doc, "信永中和审计事务所 / 实习生", "2023.11 — 2024.06", left_bold=True)
    add_bullet(doc, "参与 2 个 IPO 项目与 5 个年审项目，跨项目组对接甲方财务、律师、评估等多部门")
    add_bullet(doc, "完成穿行测试 140 项、签收单稽核 3,287 份，固定资产抽盘两百余项，审查高管银行流水 2,000+ 条")
    add_bullet(doc, "编纂审计底稿 26 份，稽核合同关键条款 508 份，上级满意度 99%")

    add_item(doc, "西安理工大学国旗护卫队 / 刀手", "2022.09 — 2024.09", left_bold=True)
    add_bullet(doc, "首批通过擎旗手与指挥刀手双项资格考核；参与校运会、换旗仪式、七一建党等活动 13 次，零失误")
    add_bullet(doc, "累计训练 60+ 次（180H+），早六点晨训出勤率 100%")

    add_item(doc, "国家税务局徐州市税务第一分局 / 暑期实习生", "2020.07 — 2020.09", left_bold=True)
    add_bullet(doc, "承包契税全部门文印工作，发现流程瓶颈并改进方法，文印效率提升近 3 倍")
    add_bullet(doc, "纳税高峰期为群众答疑引导，满意度超 90%，同批唯一优秀实习生")

    add_item(doc, "经济管理学院 / 团支书", "2018.09 — 2022.06", left_bold=True)
    add_bullet(doc, "组织团支部活动 30+ 场；策划门票制冰雪节晚会（百余人参与，学院规模最大的学生营利性晚会）")
    add_bullet(doc, "身兼班长、团支书、学习委员三职，对接辅导员与各科老师，处理文件 800+ 份，满意度 99%")

    # ═══ 教育背景 ═══
    add_section_heading(doc, "教育背景")
    add_item(doc, "西安理工大学 / 会计专业硕士", "2022.09 — 2025.06")
    add_body(doc, "主修：审计理论与实务(96)、财务会计理论与实务(90)、管理经济学(88)、大数据与会计(92)")
    add_item(doc, "南京财经大学红山学院 / 金融专业本科", "2018.09 — 2022.06")
    add_body(doc, "主修：微观经济学(98)、宏观经济学(88)")

    # ═══ 奖项荣誉 ═══
    add_section_heading(doc, "奖项荣誉")
    awards_b = [
        "2024  国旗卫士、优秀工作者、国家学业奖学金",
        "2023  校运动会团体及个人一等奖、二等学业奖学金",
        "2020  徐州市「彭聚菁英」暑期实习优秀实习生",
        "2019  校祖国 70 周年主题演讲大赛第 2 名、优秀班干部",
        "2018  校辩论赛第 3 名、优秀班干部、优秀团员、二等奖学金",
    ]
    for a in awards_b:
        add_bullet(doc, a)

    path = os.path.join(OUT_DIR, "夏樱溪-简历-AI产品经理版B.docx")
    doc.save(path)
    print(f"Version B saved: {path}")


# ═══════════════════════════════════════════
# VERSION C — 微调版（版A + 补JD缺口）
# ═══════════════════════════════════════════
def build_version_c():
    doc = Document()
    set_default_style(doc)

    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Header
    add_name_header(doc, "夏樱溪")
    add_contact_line(doc, "178 0265 8725  |  1245987829@qq.com  |  女 · 汉族 · 2001.01")

    # ═══ 自我评价 ═══
    add_section_heading(doc, "自我评价")
    self_eval = (
        "零编程基础3天从不会翻墙到独立交付全功能App并提交App Store审核，用行动证明了快速学习和从0到1的落地能力。"
        "硕士论文依靠AI工具完成8倍常规工作量实证分析并获优秀评分，具备数据驱动的分析决策意识。"
        "信永中和2个IPO+5个年审项目跨部门协调经验，善于拆解复杂流程、输出解决方案。"
        "辩论队带队校赛第三、国护队首批双考核通过，抗压自驱、沟通推动是底色。"
        "对各领域知识保持好奇并能融会贯通迁移应用。"
    )
    add_body(doc, self_eval, font_size=10)

    # ═══ 产品项目 ═══
    add_section_heading(doc, "产品项目")
    add_item(doc, "土豆ToDo — AI辅助独立开发习惯打卡App", "2025.05", left_bold=True)
    add_bullet(doc, "零编程基础自学vibe coding，3天内从不会翻墙到完整交付，已提交App Store审核")
    add_bullet(doc, "调研10+款习惯打卡类产品，识别差异化切入点，独立完成产品定位、原型设计、PRD与功能规格")
    add_bullet(doc, "全程使用Cursor/Claude辅助开发（SwiftUI + SwiftData + Widget），独立设计数据模型与App Groups共享架构")
    add_bullet(doc, "深入理解prompt engineering在真实产品开发中的边界与效率，具备AI工具二次开发与智能体搭建经验")

    # ═══ 实习经历 ═══
    add_section_heading(doc, "实习经历")

    add_item(doc, "信永中和审计事务所 / 审计实习生", "2023.11 — 2024.06", left_bold=True)
    add_bullet(doc, "参与2个IPO项目与5个年审项目，跨项目组协调对接甲方财务、律师、评估等多部门，推动项目按时交付")
    add_bullet(doc, "拆解复杂审计流程，完成穿行测试140项、签收单稽核3,287份，审查高管银行流水2,000+条")
    add_bullet(doc, "编纂审计底稿26份，稽核合同关键条款508份，上级满意度99%")

    add_item(doc, "国家税务局徐州市税务第一分局 / 暑期实习生", "2020.07 — 2020.09", left_bold=True)
    add_bullet(doc, "发现文印流程瓶颈，主动优化操作方法，部门文印效率提升近3倍")
    add_bullet(doc, "纳税高峰期直面群众需求，解答引导并维持秩序，满意度超90%，同批唯一优秀实习生")

    # ═══ 校园经历 ═══
    add_section_heading(doc, "校园经历")

    add_item(doc, "西安理工大学国旗护卫队 / 刀手", "2022.09 — 2024.09", left_bold=True)
    add_bullet(doc, "首批通过擎旗手与指挥刀手双项资格考核，参与校运会、换旗仪式等活动13次零失误，获国旗卫士称号")

    add_item(doc, "经济管理学院 / 团支书", "2018.09 — 2022.06", left_bold=True)
    add_bullet(doc, "组织团支部活动30+场；策划门票制冰雪节晚会，负责宣传曲制作、8轮线下宣讲，参与人次过百，为学院规模最大的学生营利性晚会")
    add_bullet(doc, "身兼班长、团支书、学习委员三职，对接辅导员与各科老师，处理文件800+份，满意度99%")

    # ═══ 专业技能 ═══
    add_section_heading(doc, "专业技能")
    add_bullet(doc, "AI工具：深度使用Cursor/Claude/GPT/DeepSeek，具备Prompt工程、AI工作流设计与智能体搭建经验")
    add_bullet(doc, "产品工具：熟练使用Figma进行原型设计，能独立撰写PRD与产品需求文档")
    add_bullet(doc, "数据分析：会计专业训练+论文实证分析经验，擅长从数据中提取洞察、驱动产品决策")
    add_bullet(doc, "行业认知：熟悉内容创作与社交平台运营（小红书、抖音），了解电商与营销基本逻辑")

    # ═══ 教育背景 ═══
    add_section_heading(doc, "教育背景")
    add_item(doc, "西安理工大学 / 会计专业硕士", "2022.09 — 2025.06")
    add_item(doc, "南京财经大学红山学院 / 金融专业本科", "2018.09 — 2022.06")

    # ═══ 奖项荣誉 ═══
    add_section_heading(doc, "奖项荣誉")
    awards = [
        "2024  国旗卫士、优秀工作者、国家学业奖学金",
        "2023  校运动会团体及个人一等奖、二等学业奖学金",
        "2020  徐州市「彭聚菁英」暑期实习优秀实习生",
        "2019  校祖国70周年主题演讲大赛第2名、优秀班干部",
        "2018  校辩论赛第3名、优秀班干部、优秀团员、二等奖学金",
    ]
    for a in awards:
        add_bullet(doc, a)

    path = os.path.join(OUT_DIR, "夏樱溪-简历-AI产品经理版C-微调.docx")
    doc.save(path)
    print(f"Version C saved: {path}")


# ═══════════════════════════════════════════
# VERSION D — 全面翻新版（按JD关键词重组织）
# ═══════════════════════════════════════════
def build_version_d():
    doc = Document()
    set_default_style(doc)

    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Header
    add_name_header(doc, "夏樱溪")
    add_contact_line(doc, "178 0265 8725  |  1245987829@qq.com  |  女 · 汉族 · 2001.01")

    # ═══ 自我评价 ═══
    add_section_heading(doc, "自我评价")
    self_eval = (
        "零编程基础3天独立交付全功能App并提交App Store审核——这是我能给团队带来的核心特质：缺什么就快速学会什么，不等不靠，从0到1闭环落地。"
        "硕士论文依靠AI工具完成8倍常规工作量的实证分析并获优秀评分，具备数据驱动决策的能力。"
        "信永中和2个IPO+5个年审项目经验，擅长在多部门协同中拆解复杂需求、提出可落地的解决方案。"
        "对各领域知识保持好奇心，学到的都能融会贯通并迁移应用。"
    )
    add_body(doc, self_eval, font_size=10)

    # ═══ 产品项目 ═══
    add_section_heading(doc, "产品项目")

    add_item(doc, "土豆ToDo — 习惯打卡App（0-1独立交付）", "2025.05", left_bold=True)
    add_bullet(doc, "【用户洞察】调研10+款竞品，识别「简单纯粹、无社交压力」的差异化定位，锁定轻中度习惯养成用户群")
    add_bullet(doc, "【产品设计】独立完成产品定位、信息架构、交互原型（Figma）与功能规格文档，覆盖今日打卡/日历/统计/设置4个Tab + 3个小组件")
    add_bullet(doc, "【AI辅助开发】全程使用Cursor/Claude辅助编程（SwiftUI + SwiftData），独立设计数据模型与App Groups共享架构，3天完成全流程开发并提交App Store审核")
    add_bullet(doc, "【迭代闭环】根据App Store审核反馈快速修复问题，理解产品上线后的迭代节奏与用户反馈闭环")

    # ═══ 产品核心能力 ═══
    add_section_heading(doc, "产品核心能力举证")

    # Table approach: each row is a product competency
    capabilities = [
        ("AI工具与Prompt工程",
         "硕士论文用DeepSeek/GPT/Cursor完成8倍工作量实证代码，盲审获优秀论文；独立用vibe coding完成App全流程开发，深入理解prompt engineering边界与效率；熟悉AI工作流设计与智能体搭建思路"),
        ("需求分析与PRD撰写",
         "审计项目经验训练了从复杂信息中拆解关键需求的能力；独立完成土豆ToDo的产品定位、功能规格与交互方案设计，具备从0到1的PRD撰写实战经验"),
        ("竞品分析与用户洞察",
         "产品启动阶段调研10+款习惯打卡类App并输出差异化定位；小红书内容创作在百赞账号上单条破400赞，验证了对用户爽点的捕捉能力"),
        ("数据分析与决策",
         "会计专业+8倍实证论文训练了从数据中提取洞察、量化验证假设的思维；审计经历中通过抽样测试与数据稽核定位风险点，与产品数据分析方法论高度相通"),
        ("跨团队协作推动",
         "信永中和2个IPO+5个年审项目，高频对接甲方财务、律师、评估等多部门，推动项目按节点交付；辩论队带队校赛第三、职业规划大赛院第一"),
        ("快速学习与跨界迁移",
         "3天从不会翻墙到独立交付App；一周精通高端PPT制作；市场嗅觉迁移至内容创作获400+赞；AI P图错过窗口期仍获第4名及500元奖金"),
    ]
    for title, body in capabilities:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run_t = p.add_run(title)
        run_t.bold = True
        run_t.font.size = Pt(10)
        run_t.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        add_body(doc, body)

    # ═══ 实习经历 ═══
    add_section_heading(doc, "实习经历")

    add_item(doc, "信永中和审计事务所 / 审计实习生", "2023.11 — 2024.06", left_bold=True)
    add_bullet(doc, "参与2个IPO项目与5个年审项目，在多项目组并行中协调甲方财务、律师、评估等多部门，推动项目按时交付")
    add_bullet(doc, "梳理并执行审计流程：穿行测试140项、签收单稽核3,287份、固定资产抽盘200+项、高管银行流水审查2,000+条")
    add_bullet(doc, "编纂审计底稿26份，稽核合同关键条款（验收条件、日期、付款方式等）508份，上级满意度99%")

    add_item(doc, "国家税务局徐州市税务第一分局 / 暑期实习生", "2020.07 — 2020.09", left_bold=True)
    add_bullet(doc, "发现部门文印流程瓶颈，主动优化操作方法，整体效率提升近3倍")
    add_bullet(doc, "纳税高峰期直面群众需求，解答引导并维持秩序，满意度超90%，同批唯一优秀实习生")

    # ═══ 校园经历 ═══
    add_section_heading(doc, "校园经历")

    add_item(doc, "西安理工大学国旗护卫队 / 刀手", "2022.09 — 2024.09", left_bold=True)
    add_bullet(doc, "首批通过擎旗手与指挥刀手双项考核，参与校运会、换旗仪式等活动13次零失误，获国旗卫士称号")
    add_item(doc, "经济管理学院 / 团支书", "2018.09 — 2022.06", left_bold=True)
    add_bullet(doc, "组织团支部活动30+场；策划百人冰雪节晚会（学院规模最大学生营利性晚会），负责宣传曲制作与8轮线下宣讲")
    add_bullet(doc, "身兼班长、团支书、学习委员三职，处理文件800+份，满意度99%")

    # ═══ 技能 & 工具 ═══
    add_section_heading(doc, "技能 & 工具")
    skills = [
        "AI工具：Cursor / Claude / GPT / DeepSeek — Prompt工程、AI工作流设计、智能体搭建",
        "产品工具：Figma（原型设计）、墨刀 — PRD撰写、需求文档、信息架构",
        "数据分析：Python实证分析、Excel高级应用、数据可视化",
        "行业认知：小红书/抖音平台运营、内容创作与用户增长、电商基础逻辑",
        "语言：英语（CET-4/CET-6）、普通话",
    ]
    for s in skills:
        add_bullet(doc, s)

    # ═══ 教育背景 ═══
    add_section_heading(doc, "教育背景")
    add_item(doc, "西安理工大学 / 会计专业硕士", "2022.09 — 2025.06")
    add_item(doc, "南京财经大学红山学院 / 金融专业本科", "2018.09 — 2022.06")

    # ═══ 奖项荣誉 ═══
    add_section_heading(doc, "奖项荣誉")
    awards = [
        "2024  国旗卫士、优秀工作者、国家学业奖学金",
        "2023  校运动会团体及个人一等奖、二等学业奖学金",
        "2020  徐州市「彭聚菁英」暑期实习优秀实习生",
        "2019  校祖国70周年主题演讲大赛第2名、优秀班干部",
        "2018  校辩论赛第3名、优秀班干部、优秀团员、二等奖学金",
    ]
    for a in awards:
        add_bullet(doc, a)

    path = os.path.join(OUT_DIR, "夏樱溪-简历-AI产品经理版D-全面翻新.docx")
    doc.save(path)
    print(f"Version D saved: {path}")


# ═══════════════════════════════════════════
# VERSION E — 诚实微调版（加照片 + 一页排版）
# ═══════════════════════════════════════════
def build_version_e():
    from docx.shared import Inches
    from docx.oxml.ns import qn

    doc = Document()
    set_default_style(doc)

    section = doc.sections[0]
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(0.4)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # -- Tight spacing helpers for one-page layout --
    def tight_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '999999')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def tight_item(left_text, right_text="", left_bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if left_text:
            run_l = p.add_run(left_text)
            run_l.bold = left_bold
            run_l.font.size = Pt(9.5)
            run_l.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        if right_text:
            run_r = p.add_run(f"    {right_text}")
            run_r.font.size = Pt(9)
            run_r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        return p

    def tight_bullet(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.3)
        run = p.add_run("· " + text)
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        return p

    def tight_body(text, fs=9):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.font.size = Pt(fs)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        return p

    # ═══ 表头：姓名/联系方式 + 照片 ═══
    photo_path = os.path.expanduser("~/Desktop/简历照片.pic.jpg")
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = True

    # Remove table borders
    tbl = header_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Left cell: name + contact
    left_cell = header_table.cell(0, 0)
    left_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    left_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    run_name = left_cell.paragraphs[0].add_run("夏樱溪")
    run_name.bold = True
    run_name.font.size = Pt(17)
    run_name.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    contact_p = left_cell.add_paragraph()
    contact_p.paragraph_format.space_before = Pt(2)
    contact_p.paragraph_format.space_after = Pt(0)
    run_contact = contact_p.add_run("178 0265 8725  |  1245987829@qq.com  |  女 · 汉族 · 2001.01")
    run_contact.font.size = Pt(9)
    run_contact.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Right cell: photo
    if os.path.exists(photo_path):
        right_cell = header_table.cell(0, 1)
        right_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        right_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_img = right_cell.paragraphs[0].add_run()
        run_img.add_picture(photo_path, width=Inches(0.72))

    # Thin gray line after header
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(2)
    p_sep.paragraph_format.space_after = Pt(0)
    pPr = p_sep._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ═══ 自我评价 ═══
    tight_section_heading("自我评价")
    tight_body(
        "零编程基础自学vibe coding，3天内从不会翻墙到独立交付全功能App并提交App Store审核——"
        "缺什么就快速学会什么，不等不靠。硕士论文用AI工具完成8倍常规工作量的实证分析并获优秀评分。"
        "信永中和2个IPO+5个年审项目跨部门协调经验，善于拆解复杂流程、推动多方协作落地。"
        "辩论队带队校赛第三、国护队首批双考核通过，抗压自驱、沟通推动是底色。"
        "对各领域知识保持好奇并能融会贯通迁移应用。"
    )

    # ═══ 产品项目 ═══
    tight_section_heading("产品项目")
    tight_item("土豆ToDo — AI辅助独立开发习惯打卡App", "2025.05", left_bold=True)
    tight_bullet("调研10+款习惯打卡类产品，识别「简单纯粹、无社交压力」的差异化定位，确定核心功能方向")
    tight_bullet("全程使用Cursor/Claude通过自然语言描述完成产品交互方案、界面设计与功能开发，3天内完整交付")
    tight_bullet("独立完成App数据结构定义（习惯、打卡记录、倒计时事件等实体）、4个Tab信息架构与App Groups小组件共享方案")
    tight_bullet("经历App Store提审与反馈修复流程，理解产品上线迭代的基本节奏")

    # ═══ 实习经历 ═══
    tight_section_heading("实习经历")

    tight_item("信永中和审计事务所 / 审计实习生", "2023.11 — 2024.06", left_bold=True)
    tight_bullet("参与2个IPO项目与5个年审项目，跨项目组协调对接甲方财务、律师、评估等多部门，推动项目按时交付")
    tight_bullet("拆解复杂审计流程，完成穿行测试140项、签收单稽核3,287份，审查高管银行流水2,000+条")
    tight_bullet("编纂审计底稿26份，稽核合同关键条款508份，上级满意度99%")

    tight_item("国家税务局徐州市税务第一分局 / 暑期实习生", "2020.07 — 2020.09", left_bold=True)
    tight_bullet("发现文印流程瓶颈，主动优化操作方法，部门文印效率提升近3倍")
    tight_bullet("纳税高峰期直面群众需求，解答引导并维持秩序，满意度超90%，同批唯一优秀实习生")

    # ═══ 校园经历 ═══
    tight_section_heading("校园经历")

    tight_item("西安理工大学国旗护卫队 / 刀手", "2022.09 — 2024.09", left_bold=True)
    tight_bullet("首批通过擎旗手与指挥刀手双项资格考核，参与校运会、换旗仪式等活动13次零失误，获国旗卫士称号")

    tight_item("经济管理学院 / 团支书", "2018.09 — 2022.06", left_bold=True)
    tight_bullet("组织团支部活动30+场；策划门票制冰雪节晚会，负责宣传曲制作、8轮线下宣讲，参与人次过百，为学院规模最大的学生营利性晚会")
    tight_bullet("身兼班长、团支书、学习委员三职，对接辅导员与各科老师，处理文件800+份，满意度99%")

    # ═══ 技能掌握 ═══
    tight_section_heading("技能掌握")
    tight_bullet("AI工具：深度使用Cursor/Claude/GPT/DeepSeek辅助编程与内容生成，熟悉Prompt工程在实际开发中的调试方法")
    tight_bullet("产品思维：善用自然语言描述需求并与AI协作产出交互方案，具备产品功能定位与信息架构设计能力")
    tight_bullet("数据分析：会计专业+硕士论文实证分析（Python），擅长从数据中提取洞察以支撑决策")
    tight_bullet("行业认知：小红书/抖音内容创作经验，了解平台运营与用户增长基本逻辑；高端PPT制作能力")

    # ═══ 教育背景 ═══
    tight_section_heading("教育背景")
    tight_item("西安理工大学 / 会计专业硕士", "2022.09 — 2025.06")
    tight_item("南京财经大学红山学院 / 金融专业本科", "2018.09 — 2022.06")

    # ═══ 奖项荣誉 ═══
    tight_section_heading("奖项荣誉")
    tight_body(
        "2024 国旗卫士、优秀工作者、国家学业奖学金  |  "
        "2023 校运动会一等奖、二等学业奖学金  |  "
        "2020 徐州市「彭聚菁英」优秀实习生  |  "
        "2019 校演讲大赛第2名、优秀班干部  |  "
        "2018 校辩论赛第3名、优秀班干部、优秀团员",
        fs=7.5
    )

    path = os.path.join(OUT_DIR, "夏樱溪-简历-版E.docx")
    doc.save(path)
    print(f"Version E saved: {path}")


# ═══════════════════════════════════════════
# VERSION F — 诚实全面翻新版（只写真的）
# ═══════════════════════════════════════════
def build_version_f():
    doc = Document()
    set_default_style(doc)

    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    add_name_header(doc, "夏樱溪")
    add_contact_line(doc, "178 0265 8725  |  1245987829@qq.com  |  女 · 汉族 · 2001.01")

    # ═══ 自我评价 ═══
    add_section_heading(doc, "自我评价")
    self_eval = (
        "零编程基础自学vibe coding，3天从不会翻墙到独立交付全功能App并提交App Store审核——"
        "这是我能给团队带来的核心特质：缺什么就快速学会什么，不等不靠，从0到1闭环落地。"
        "硕士论文依靠AI工具完成8倍常规工作量的实证分析并获优秀评分。"
        "信永中和2个IPO+5个年审项目经验，擅长在多部门协同中拆解复杂需求、推动落地。"
        "对各领域知识保持好奇心，学到的都能融会贯通并迁移应用。"
    )
    add_body(doc, self_eval, font_size=10)

    # ═══ 产品项目 ═══
    add_section_heading(doc, "产品项目")

    add_item(doc, "土豆ToDo — 习惯打卡App（0-1独立交付）", "2025.05", left_bold=True)
    add_bullet(doc, "【用户洞察】调研10+款习惯打卡类竞品，识别「简单纯粹、无社交压力」的差异化定位，锁定轻中度习惯养成用户群")
    add_bullet(doc, "【产品设计】通过自然语言描述与AI协作，独立完成产品定位、信息架构（4个Tab+3个小组件）与交互方案设计")
    add_bullet(doc, "【数据建模】定义App核心数据结构——习惯、打卡记录、倒计时事件、番茄钟会话等实体及其关联关系，设计App Groups共享方案供Widget读取")
    add_bullet(doc, "【AI辅助开发】全程使用Cursor/Claude辅助编程（SwiftUI + SwiftData），3天内从环境搭建到完整交付")
    add_bullet(doc, "【迭代闭环】经历App Store提审、被拒、修复的完整反馈循环，理解产品上线后的迭代节奏与用户反馈闭环")

    # ═══ 产品相关能力 ═══
    add_section_heading(doc, "产品相关能力举证")

    capabilities = [
        ("AI工具深度使用",
         "硕士论文用DeepSeek/GPT/Cursor辅助完成8倍常规工作量的实证代码，校内盲审获优秀论文评分；独立用vibe coding完成App全流程开发，对prompt engineering在真实开发中的调试和效率边界有直接体感"),
        ("用户洞察与需求挖掘",
         "产品启动阶段调研10+款竞品并输出差异化定位；小红书内容创作在平均百赞账号上单条破400赞，验证了对用户喜好的捕捉能力；审计工作中与甲方多部门高频沟通，训练了从复杂信息中提取核心需求的能力"),
        ("逻辑拆解与方案设计",
         "审计项目训练了将复杂业务流程拆解为可执行步骤的能力（穿行测试140项、底稿26份）；独立完成App从功能定位到数据结构的完整方案设计"),
        ("跨团队协作与推动",
         "信永中和2个IPO+5个年审项目，高频对接甲方财务、律师、评估等多部门，推动项目按节点交付；辩论队带队校赛第三，职业规划大赛院第一"),
        ("快速学习与迁移应用",
         "3天从不会翻墙到独立交付App；一周精通高端PPT制作；内容嗅觉迁移至创作——单条破400赞；AI P图错过流量窗口期仍获第4名及500元奖金"),
        ("数据驱动的分析意识",
         "会计专业核心训练+8倍实证论文：从数据中定位问题、量化验证假设；审计经历中通过抽样与数据稽核识别风险——此方法论与产品数据分析高度相通"),
    ]
    for title, body in capabilities:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run_t = p.add_run(title)
        run_t.bold = True
        run_t.font.size = Pt(10)
        run_t.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        add_body(doc, body)

    # ═══ 实习经历 ═══
    add_section_heading(doc, "实习经历")

    add_item(doc, "信永中和审计事务所 / 审计实习生", "2023.11 — 2024.06", left_bold=True)
    add_bullet(doc, "参与2个IPO项目与5个年审项目，在多项目组并行中协调甲方财务、律师、评估等多部门，推动项目按时交付")
    add_bullet(doc, "梳理并执行审计流程：穿行测试140项、签收单稽核3,287份、固定资产抽盘200+项、高管银行流水审查2,000+条")
    add_bullet(doc, "编纂审计底稿26份，稽核合同关键条款（验收条件、日期、付款方式等）508份，上级满意度99%")

    add_item(doc, "国家税务局徐州市税务第一分局 / 暑期实习生", "2020.07 — 2020.09", left_bold=True)
    add_bullet(doc, "发现部门文印流程瓶颈，主动优化操作方法，整体效率提升近3倍")
    add_bullet(doc, "纳税高峰期直面群众需求，解答引导并维持秩序，满意度超90%，同批唯一优秀实习生")

    # ═══ 校园经历 ═══
    add_section_heading(doc, "校园经历")

    add_item(doc, "西安理工大学国旗护卫队 / 刀手", "2022.09 — 2024.09", left_bold=True)
    add_bullet(doc, "首批通过擎旗手与指挥刀手双项考核，参与校运会、换旗仪式等活动13次零失误，获国旗卫士称号")
    add_item(doc, "经济管理学院 / 团支书", "2018.09 — 2022.06", left_bold=True)
    add_bullet(doc, "组织团支部活动30+场；策划百人冰雪节晚会（学院规模最大学生营利性晚会），负责宣传曲制作与8轮线下宣讲")
    add_bullet(doc, "身兼班长、团支书、学习委员三职，处理文件800+份，满意度99%")

    # ═══ 工具 & 技能 ═══
    add_section_heading(doc, "工具 & 技能")
    skills = [
        "AI编程：Cursor / Claude — 通过自然语言驱动AI完成SwiftUI应用开发，能在面试中demo全部开发过程",
        "AI对话：GPT / DeepSeek — 熟练用于论文实证代码生成、内容创作、竞品分析等场景",
        "数据分析：Python（论文实证）、Excel高级应用",
        "内容平台：小红书、抖音内容创作与运营经验，能理解平台规则与用户偏好",
        "其他：高端PPT制作、基础音视频剪辑",
        "语言：英语四级",
    ]
    for s in skills:
        add_bullet(doc, s)

    # ═══ 教育背景 ═══
    add_section_heading(doc, "教育背景")
    add_item(doc, "西安理工大学 / 会计专业硕士", "2022.09 — 2025.06")
    add_item(doc, "南京财经大学红山学院 / 金融专业本科", "2018.09 — 2022.06")

    # ═══ 奖项荣誉 ═══
    add_section_heading(doc, "奖项荣誉")
    awards = [
        "2024  国旗卫士、优秀工作者、国家学业奖学金",
        "2023  校运动会团体及个人一等奖、二等学业奖学金",
        "2020  徐州市「彭聚菁英」暑期实习优秀实习生",
        "2019  校祖国70周年主题演讲大赛第2名、优秀班干部",
        "2018  校辩论赛第3名、优秀班干部、优秀团员、二等奖学金",
    ]
    for a in awards:
        add_bullet(doc, a)

    path = os.path.join(OUT_DIR, "夏樱溪-简历-版F-诚实全面翻新.docx")
    doc.save(path)
    print(f"Version F saved: {path}")


# ═══════════════════════════════════════════
# Layout helpers (reusable across G/H/I)
# ═══════════════════════════════════════════

LAYOUT_CONTENT = {
    "self_eval": (
        "零编程基础自学vibe coding，3天内从不会翻墙到独立交付全功能App并提交App Store审核——"
        "缺什么就快速学会什么，不等不靠。硕士论文用AI工具完成8倍常规工作量的实证分析并获优秀评分。"
        "信永中和2个IPO+5个年审项目跨部门协调经验，善于拆解复杂流程、推动多方协作落地。"
        "辩论队带队校赛第三、国护队首批双考核通过，抗压自驱、沟通推动是底色。"
        "对各领域知识保持好奇并能融会贯通迁移应用。"
    ),
    "skills": [
        "AI工具：深度使用Cursor/Claude/GPT/DeepSeek辅助编程与内容生成，熟悉Prompt工程在实际开发中的调试方法",
        "产品思维：善用自然语言描述需求并与AI协作产出交互方案，具备产品功能定位与信息架构设计能力",
        "数据分析：会计专业+硕士论文实证分析（Python），擅长从数据中提取洞察以支撑决策",
        "行业认知：小红书/抖音内容创作经验，了解平台运营与用户增长基本逻辑；高端PPT制作能力",
    ],
    "awards_line": (
        "2024 国旗卫士、优秀工作者、国家学业奖学金 | 2023 校运动会一等奖、二等学业奖学金 | "
        "2020 徐州市「彭聚菁英」优秀实习生 | 2019 校演讲大赛第2名、优秀班干部 | "
        "2018 校辩论赛第3名、优秀班干部、优秀团员"
    ),
}

def _add_text_run(paragraph, text, bold=False, size=10, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return run

def _make_borderless_table(doc, rows, cols):
    table = doc.add_table(rows=rows, cols=cols)
    table.autofit = True
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for bn in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{bn}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)
    return table

def _clear_cell_para(cell, idx=0):
    """Clear and return a clean paragraph from a cell."""
    p = cell.paragraphs[idx]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p

def _add_sep_line(doc, color="CCCCCC"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def _section_head(doc, text, fs=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(fs)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def _item(doc, left, right="", left_bold=False, left_fs=10, right_fs=9):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if left:
        r = p.add_run(left)
        r.bold = left_bold
        r.font.size = Pt(left_fs)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    if right:
        r2 = p.add_run(f"    {right}")
        r2.font.size = Pt(right_fs)
        r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    return p

def _bullet(doc, text, fs=9):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    r = p.add_run("· " + text)
    r.font.size = Pt(fs)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p

def _body(doc, text, fs=9.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    r.font.size = Pt(fs)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p


# ═══════════════════════════════════════════
# VERSION G — 左窄右宽侧边栏布局
# ═══════════════════════════════════════════
def build_version_g():
    doc = Document()
    set_default_style(doc)

    section = doc.sections[0]
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(0.6)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    photo_path = os.path.expanduser("~/Desktop/简历照片.pic.jpg")

    # Main 2-column wrapper
    main = _make_borderless_table(doc, 1, 2)

    # === LEFT COLUMN (33%) ===
    left_cell = main.cell(0, 0)
    left_cell.width = Cm(5.5)

    # Photo
    if os.path.exists(photo_path):
        p_img = _clear_cell_para(left_cell, 0)
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        run_img.add_picture(photo_path, width=Cm(3.2))

    # Name
    pn = left_cell.add_paragraph()
    pn.paragraph_format.space_before = Pt(4)
    pn.paragraph_format.space_after = Pt(0)
    pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_text_run(pn, "夏樱溪", bold=True, size=16, color=RGBColor(0x22, 0x22, 0x22))

    # Contact info stacked
    for line in ["178 0265 8725", "1245987829@qq.com", "女 · 汉族 · 2001.01"]:
        pc = left_cell.add_paragraph()
        pc.paragraph_format.space_before = Pt(1)
        pc.paragraph_format.space_after = Pt(1)
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_text_run(pc, line, size=8.5, color=RGBColor(0x88, 0x88, 0x88))

    _add_sep_line(left_cell)

    # Skills in sidebar
    p_sk_h = left_cell.add_paragraph()
    p_sk_h.paragraph_format.space_before = Pt(6)
    p_sk_h.paragraph_format.space_after = Pt(2)
    _add_text_run(p_sk_h, "技能掌握", bold=True, size=10)
    _add_sep_line(left_cell, color="DDDDDD")

    for s in LAYOUT_CONTENT["skills"]:
        p_sk = left_cell.add_paragraph()
        p_sk.paragraph_format.space_before = Pt(1)
        p_sk.paragraph_format.space_after = Pt(1)
        _add_text_run(p_sk, "· " + s, size=8, color=RGBColor(0x55, 0x55, 0x55))

    # Education in sidebar
    p_ed_h = left_cell.add_paragraph()
    p_ed_h.paragraph_format.space_before = Pt(8)
    p_ed_h.paragraph_format.space_after = Pt(2)
    _add_text_run(p_ed_h, "教育背景", bold=True, size=10)
    _add_sep_line(left_cell, color="DDDDDD")

    for edu_text in [
        "西安理工大学 / 会计专业硕士",
        "2022.09 — 2025.06",
        "南京财经大学红山学院 / 金融本科",
        "2018.09 — 2022.06",
    ]:
        pe = left_cell.add_paragraph()
        pe.paragraph_format.space_before = Pt(0)
        pe.paragraph_format.space_after = Pt(0)
        _add_text_run(pe, edu_text, size=8, color=RGBColor(0x55, 0x55, 0x55))

    # Awards in sidebar
    p_aw_h = left_cell.add_paragraph()
    p_aw_h.paragraph_format.space_before = Pt(8)
    p_aw_h.paragraph_format.space_after = Pt(2)
    _add_text_run(p_aw_h, "奖项荣誉", bold=True, size=10)
    _add_sep_line(left_cell, color="DDDDDD")

    pa = left_cell.add_paragraph()
    pa.paragraph_format.space_before = Pt(2)
    _add_text_run(pa, LAYOUT_CONTENT["awards_line"], size=7, color=RGBColor(0x66, 0x66, 0x66))

    # === RIGHT COLUMN (67%) ===
    right_cell = main.cell(0, 1)
    right_cell.width = Cm(11)

    # Self eval
    _section_head(right_cell, "自我评价", fs=10.5)
    _body(right_cell, LAYOUT_CONTENT["self_eval"], fs=9)

    # Product project
    _section_head(right_cell, "产品项目", fs=10.5)
    _item(right_cell, "土豆ToDo — AI辅助独立开发习惯打卡App", "2025.05", left_bold=True)
    _bullet(right_cell, "调研10+款习惯打卡类产品，识别「简单纯粹、无社交压力」的差异化定位，确定核心功能方向", fs=8.5)
    _bullet(right_cell, "全程使用Cursor/Claude通过自然语言描述完成产品交互方案、界面设计与功能开发，3天内完整交付", fs=8.5)
    _bullet(right_cell, "独立完成App数据结构定义（习惯、打卡记录、倒计时事件等实体）、4个Tab信息架构与App Groups小组件共享方案", fs=8.5)
    _bullet(right_cell, "经历App Store提审与反馈修复流程，理解产品上线迭代的基本节奏", fs=8.5)

    # Experience
    _section_head(right_cell, "实习经历", fs=10.5)

    _item(right_cell, "信永中和审计事务所 / 审计实习生", "2023.11 — 2024.06", left_bold=True)
    _bullet(right_cell, "参与2个IPO项目与5个年审项目，跨项目组协调对接甲方财务、律师、评估等多部门，推动项目按时交付", fs=8.5)
    _bullet(right_cell, "拆解复杂审计流程，完成穿行测试140项、签收单稽核3,287份，审查高管银行流水2,000+条", fs=8.5)
    _bullet(right_cell, "编纂审计底稿26份，稽核合同关键条款508份，上级满意度99%", fs=8.5)

    _item(right_cell, "国家税务局徐州市税务第一分局 / 暑期实习生", "2020.07 — 2020.09", left_bold=True)
    _bullet(right_cell, "发现文印流程瓶颈，主动优化操作方法，部门文印效率提升近3倍", fs=8.5)
    _bullet(right_cell, "纳税高峰期直面群众需求，解答引导并维持秩序，满意度超90%，同批唯一优秀实习生", fs=8.5)

    # Campus
    _section_head(right_cell, "校园经历", fs=10.5)
    _item(right_cell, "西安理工大学国旗护卫队 / 刀手", "2022.09 — 2024.09", left_bold=True)
    _bullet(right_cell, "首批通过擎旗手与指挥刀手双项资格考核，参与校运会、换旗仪式等活动13次零失误，获国旗卫士称号", fs=8.5)
    _item(right_cell, "经济管理学院 / 团支书", "2018.09 — 2022.06", left_bold=True)
    _bullet(right_cell, "组织团支部活动30+场；策划门票制冰雪节晚会（百余人参与，学院规模最大的学生营利性晚会），负责宣传曲制作与8轮线下宣讲", fs=8.5)
    _bullet(right_cell, "身兼班长、团支书、学习委员三职，对接辅导员与各科老师，处理文件800+份，满意度99%", fs=8.5)

    path = os.path.join(OUT_DIR, "夏樱溪-简历-版G-左窄右宽.docx")
    doc.save(path)
    print(f"Version G saved: {path}")


# ═══════════════════════════════════════════
# VERSION H — 单栏紧凑 + 照片行内 + 底部双排
# ═══════════════════════════════════════════
def build_version_h():
    doc = Document()
    set_default_style(doc)

    section = doc.sections[0]
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    photo_path = os.path.expanduser("~/Desktop/简历照片.pic.jpg")

    # Header: photo inline with name
    header_tbl = _make_borderless_table(doc, 1, 2)
    h_left = header_tbl.cell(0, 0)
    h_left.width = Cm(2.5)
    if os.path.exists(photo_path):
        p_img = _clear_cell_para(h_left, 0)
        run_img = p_img.add_run()
        run_img.add_picture(photo_path, width=Cm(2.8))

    h_right = header_tbl.cell(0, 1)
    _clear_cell_para(h_right, 0)
    _add_text_run(h_right.paragraphs[0], "夏樱溪", bold=True, size=20, color=RGBColor(0x22, 0x22, 0x22))
    pc = h_right.add_paragraph()
    pc.paragraph_format.space_before = Pt(3)
    _add_text_run(pc, "178 0265 8725  |  1245987829@qq.com  |  女 · 汉族 · 2001.01", size=9.5, color=RGBColor(0x88, 0x88, 0x88))

    _add_sep_line(doc)

    # Self eval
    _section_head(doc, "自我评价", fs=11)
    _body(doc, LAYOUT_CONTENT["self_eval"], fs=9.5)

    # Product
    _section_head(doc, "产品项目", fs=11)
    _item(doc, "土豆ToDo — AI辅助独立开发习惯打卡App", "2025.05", left_bold=True)
    _bullet(doc, "调研10+款习惯打卡类产品，识别「简单纯粹、无社交压力」的差异化定位，确定核心功能方向", fs=9)
    _bullet(doc, "全程使用Cursor/Claude通过自然语言描述完成产品交互方案、界面设计与功能开发，3天内完整交付", fs=9)
    _bullet(doc, "独立完成App数据结构定义（习惯、打卡记录、倒计时事件等实体）、4个Tab信息架构与App Groups小组件共享方案", fs=9)
    _bullet(doc, "经历App Store提审与反馈修复流程，理解产品上线迭代的基本节奏", fs=9)

    # Experience
    _section_head(doc, "实习经历", fs=11)
    _item(doc, "信永中和审计事务所 / 审计实习生", "2023.11 — 2024.06", left_bold=True)
    _bullet(doc, "参与2个IPO项目与5个年审项目，跨项目组协调对接甲方财务、律师、评估等多部门，推动项目按时交付", fs=9)
    _bullet(doc, "拆解复杂审计流程，完成穿行测试140项、签收单稽核3,287份，审查高管银行流水2,000+条", fs=9)
    _bullet(doc, "编纂审计底稿26份，稽核合同关键条款508份，上级满意度99%", fs=9)

    _item(doc, "国家税务局徐州市税务第一分局 / 暑期实习生", "2020.07 — 2020.09", left_bold=True)
    _bullet(doc, "发现文印流程瓶颈，主动优化操作方法，部门文印效率提升近3倍", fs=9)
    _bullet(doc, "纳税高峰期直面群众需求，解答引导并维持秩序，满意度超90%，同批唯一优秀实习生", fs=9)

    # Campus
    _section_head(doc, "校园经历", fs=11)
    _item(doc, "西安理工大学国旗护卫队 / 刀手", "2022.09 — 2024.09", left_bold=True)
    _bullet(doc, "首批通过擎旗手与指挥刀手双项资格考核，参与校运会、换旗仪式等活动13次零失误，获国旗卫士称号", fs=9)
    _item(doc, "经济管理学院 / 团支书", "2018.09 — 2022.06", left_bold=True)
    _bullet(doc, "组织团支部活动30+场；策划门票制冰雪节晚会（百余人参与，学院规模最大的学生营利性晚会），负责宣传曲制作与8轮线下宣讲", fs=9)
    _bullet(doc, "身兼班长、团支书、学习委员三职，对接辅导员与各科老师，处理文件800+份，满意度99%", fs=9)

    # Skills (compact inline)
    _section_head(doc, "技能掌握", fs=11)
    for s in LAYOUT_CONTENT["skills"]:
        _bullet(doc, s, fs=9)

    # Bottom: education + awards side by side
    bot_tbl = _make_borderless_table(doc, 1, 2)
    b_left = bot_tbl.cell(0, 0)
    b_left.width = Cm(8)
    _section_head(b_left, "教育背景", fs=10.5)
    _item(b_left, "西安理工大学 / 会计专业硕士", "2022.09 — 2025.06", left_fs=9.5, right_fs=8.5)
    _item(b_left, "南京财经大学红山学院 / 金融专业本科", "2018.09 — 2022.06", left_fs=9.5, right_fs=8.5)

    b_right = bot_tbl.cell(0, 1)
    b_right.width = Cm(8.5)
    _section_head(b_right, "奖项荣誉", fs=10.5)
    pa = b_right.add_paragraph()
    pa.paragraph_format.space_before = Pt(2)
    _add_text_run(pa, LAYOUT_CONTENT["awards_line"], size=7.5, color=RGBColor(0x55, 0x55, 0x55))

    path = os.path.join(OUT_DIR, "夏樱溪-简历-版H-单栏紧凑.docx")
    doc.save(path)
    print(f"Version H saved: {path}")


# ═══════════════════════════════════════════
# VERSION I — 三段式混合分栏布局
# ═══════════════════════════════════════════
def build_version_i():
    doc = Document()
    set_default_style(doc)

    section = doc.sections[0]
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.4)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    photo_path = os.path.expanduser("~/Desktop/简历照片.pic.jpg")

    # === TOP: Business card header ===
    top_tbl = _make_borderless_table(doc, 1, 3)
    # Col 0: photo
    t0 = top_tbl.cell(0, 0)
    t0.width = Cm(2.8)
    if os.path.exists(photo_path):
        p_img = _clear_cell_para(t0, 0)
        run_img = p_img.add_run()
        run_img.add_picture(photo_path, width=Cm(2.8))

    # Col 1: name + contact
    t1 = top_tbl.cell(0, 1)
    t1.width = Cm(8.5)
    _clear_cell_para(t1, 0)
    _add_text_run(t1.paragraphs[0], "夏樱溪", bold=True, size=20, color=RGBColor(0x22, 0x22, 0x22))
    pc1 = t1.add_paragraph()
    pc1.paragraph_format.space_before = Pt(3)
    _add_text_run(pc1, "178 0265 8725  |  1245987829@qq.com  |  女 · 汉族 · 2001.01", size=9.5, color=RGBColor(0x88, 0x88, 0x88))

    # Col 2: 3-line quick tag
    t2 = top_tbl.cell(0, 2)
    t2.width = Cm(5.5)
    _clear_cell_para(t2, 0)
    t2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for tag in ["AI工具深度使用者", "0-1产品独立交付", "快速学习 · 跨界迁移"]:
        pt = t2.add_paragraph()
        pt.paragraph_format.space_before = Pt(1)
        pt.paragraph_format.space_after = Pt(1)
        pt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_text_run(pt, tag, bold=True, size=8.5, color=RGBColor(0xFF, 0xD6, 0x0A))

    # Divider
    _add_sep_line(doc)

    # === MIDDLE: 2-column content ===
    mid_tbl = _make_borderless_table(doc, 1, 2)

    # Left: self-eval + skills
    ml = mid_tbl.cell(0, 0)
    ml.width = Cm(8.3)
    _section_head(ml, "自我评价", fs=10)
    _body(ml, LAYOUT_CONTENT["self_eval"], fs=8.5)
    _section_head(ml, "技能掌握", fs=10)
    for s in LAYOUT_CONTENT["skills"]:
        _bullet(ml, s, fs=8)

    # Right: product + campus
    mr = mid_tbl.cell(0, 1)
    mr.width = Cm(8.3)
    _section_head(mr, "产品项目", fs=10)
    _item(mr, "土豆ToDo — AI辅助独立开发习惯打卡App", "2025.05", left_bold=True, left_fs=9, right_fs=8)
    _bullet(mr, "调研10+款习惯打卡类产品，识别「简单纯粹、无社交压力」的差异化定位，确定核心功能方向", fs=8)
    _bullet(mr, "全程使用Cursor/Claude通过自然语言描述完成产品交互方案、界面设计与功能开发，3天内完整交付", fs=8)
    _bullet(mr, "独立完成App数据结构定义（习惯、打卡记录、倒计时事件等实体）、4个Tab信息架构与App Groups小组件共享方案", fs=8)
    _bullet(mr, "经历App Store提审与反馈修复流程，理解产品上线迭代的基本节奏", fs=8)

    _section_head(mr, "校园经历", fs=10)
    _item(mr, "西安理工大学国旗护卫队 / 刀手", "2022.09 — 2024.09", left_bold=True, left_fs=9, right_fs=8)
    _bullet(mr, "首批通过擎旗手与指挥刀手双项资格考核，参与校运会、换旗仪式等活动13次零失误，获国旗卫士称号", fs=8)
    _item(mr, "经济管理学院 / 团支书", "2018.09 — 2022.06", left_bold=True, left_fs=9, right_fs=8)
    _bullet(mr, "组织团支部活动30+场；策划门票制冰雪节晚会（百余人参与，学院规模最大的学生营利性晚会），负责宣传曲制作与8轮线下宣讲", fs=8)
    _bullet(mr, "身兼班长、团支书、学习委员三职，处理文件800+份，满意度99%", fs=8)

    # Experience: full width below
    _section_head(doc, "实习经历", fs=10.5)

    exp_tbl = _make_borderless_table(doc, 1, 2)
    el_cell = exp_tbl.cell(0, 0)
    el_cell.width = Cm(8.3)
    _item(el_cell, "信永中和审计事务所 / 审计实习生", "2023.11 — 2024.06", left_bold=True, left_fs=9.5, right_fs=8.5)
    _bullet(el_cell, "参与2个IPO项目与5个年审项目，跨项目组协调对接甲方财务、律师、评估等多部门，推动项目按时交付", fs=8.5)
    _bullet(el_cell, "拆解复杂审计流程，完成穿行测试140项、签收单稽核3,287份，审查高管银行流水2,000+条", fs=8.5)
    _bullet(el_cell, "编纂审计底稿26份，稽核合同关键条款508份，上级满意度99%", fs=8.5)

    er_cell = exp_tbl.cell(0, 1)
    er_cell.width = Cm(8.3)
    _item(er_cell, "国家税务局徐州市税务第一分局 / 暑期实习生", "2020.07 — 2020.09", left_bold=True, left_fs=9.5, right_fs=8.5)
    _bullet(er_cell, "发现文印流程瓶颈，主动优化操作方法，部门文印效率提升近3倍", fs=8.5)
    _bullet(er_cell, "纳税高峰期直面群众需求，解答引导并维持秩序，满意度超90%，同批唯一优秀实习生", fs=8.5)

    # === BOTTOM: education + awards one line ===
    _add_sep_line(doc)

    bot_tbl = _make_borderless_table(doc, 1, 2)
    bl = bot_tbl.cell(0, 0)
    bl.width = Cm(8.3)
    pe1 = bl.add_paragraph()
    pe1.paragraph_format.space_before = Pt(2)
    pe1.paragraph_format.space_after = Pt(0)
    _add_text_run(pe1, "教育背景  |  西安理工大学 / 会计专业硕士（2022.09 — 2025.06）  |  南京财经大学红山学院 / 金融本科（2018.09 — 2022.06）", size=8, color=RGBColor(0x66, 0x66, 0x66))

    br = bot_tbl.cell(0, 1)
    br.width = Cm(8.3)
    pe2 = br.add_paragraph()
    pe2.paragraph_format.space_before = Pt(2)
    pe2.paragraph_format.space_after = Pt(0)
    pe2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_text_run(pe2, "奖项  " + LAYOUT_CONTENT["awards_line"], size=7, color=RGBColor(0x77, 0x77, 0x77))

    path = os.path.join(OUT_DIR, "夏樱溪-简历-版I-三段混合.docx")
    doc.save(path)
    print(f"Version I saved: {path}")


if __name__ == "__main__":
    build_version_g()
    build_version_h()
    build_version_i()
    print("Done!")
